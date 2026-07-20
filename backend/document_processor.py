"""
document_processor.py
----------------------
Handles turning raw input (PDF, DOCX, TXT, or a web URL) into clean text,
then splitting that text into overlapping chunks suitable for embedding
and retrieval.
"""

import re
import io
from typing import List

import pdfplumber
import docx
import requests
from bs4 import BeautifulSoup


class DocumentProcessorError(Exception):
    """Raised when a document cannot be read or parsed."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
    except Exception as exc:
        raise DocumentProcessorError(f"Could not read PDF: {exc}") from exc
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text out of any tables in the document
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as exc:
        raise DocumentProcessorError(f"Could not read DOCX: {exc}") from exc


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentProcessorError("Could not decode text file (unsupported encoding).")


def extract_text_from_url(url: str) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0 (DocSummarizer/1.0)"}
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
    except requests.RequestException as exc:
        raise DocumentProcessorError(f"Could not fetch URL: {exc}") from exc

    soup = BeautifulSoup(response.text, "html.parser")

    # Strip elements that never contain real article content
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    # Collapse excess blank lines left behind by stripped tags
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = 220, overlap: int = 40) -> List[str]:
    """
    Splits text into overlapping word-based chunks.

    chunk_size and overlap are measured in words, which keeps chunks a
    predictable size for the embedding model regardless of sentence length.
    """
    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        if end == len(words):
            break
        start = end - overlap  # step forward, keeping some overlap for context continuity
    return chunks


def process_document(source_type: str, raw: bytes | str) -> tuple[List[str], str]:
    """
    source_type: one of "pdf", "docx", "txt", "url"
    raw: file bytes for pdf/docx/txt, or the URL string for "url"
    Returns (chunks, full_cleaned_text).
    """
    if source_type == "pdf":
        text = extract_text_from_pdf(raw)
    elif source_type == "docx":
        text = extract_text_from_docx(raw)
    elif source_type == "txt":
        text = extract_text_from_txt(raw)
    elif source_type == "url":
        text = extract_text_from_url(raw)
    elif source_type == "text":
        text = raw  # already plain text, pasted directly by the user
    else:
        raise DocumentProcessorError(f"Unsupported source type: {source_type}")

    text = clean_text(text)
    if not text:
        raise DocumentProcessorError("No extractable text was found in this document.")

    return chunk_text(text), text
